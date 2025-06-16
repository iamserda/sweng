from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT


# Helper to format heading text
def add_heading(doc, text, level=1):
    doc.add_heading(text, level=level)


# Helper to format normal paragraph text
def add_paragraph(doc, text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(11)
    if bold:
        run.bold = True
    return p


# Create resumes
resumes = []

# Resume 1: 1-3 Years Experience
doc1 = Document()
add_heading(doc1, "JANE DOE", 0)
add_paragraph(
    doc1,
    "📍 San Francisco, CA | ✉️ jane.doe@gmail.com | 📞 (123) 456-7890 | 🔗 github.com/janedoe",
)
add_heading(doc1, "SUMMARY", 1)
add_paragraph(
    doc1,
    "Self-driven Backend Engineer with 2 years of experience building scalable web services and RESTful APIs using Node.js, Express, and MongoDB...",
)
add_heading(doc1, "TECHNICAL SKILLS", 1)
add_paragraph(
    doc1,
    "- Languages: JavaScript (Node.js), Python\n- Frameworks: Express.js, Flask\n- Databases: MongoDB, PostgreSQL\n- Tools: Git, Docker, Postman, Jenkins\n- Cloud: AWS (EC2, S3), Heroku\n- Testing: Jest, Mocha, PyTest",
)
add_heading(doc1, "PROFESSIONAL EXPERIENCE", 1)
add_paragraph(
    doc1,
    "Backend Developer, TechNova Inc., Jun 2023 – Present\n- Built and maintained RESTful APIs...\n",
)
add_paragraph(
    doc1,
    "Software Engineer Intern, CodeCrush Labs, Jan 2022 – May 2023\n- Collaborated with 3 developers...",
)
add_heading(doc1, "PROJECTS", 1)
add_paragraph(doc1, "TaskTrackr - A Trello-like productivity web app...")
add_heading(doc1, "EDUCATION", 1)
add_paragraph(doc1, "B.S. in Computer Science, University of California, Davis – 2021")
resumes.append(("Backend_Engineer_1-3_Years.docx", doc1))

# Resume 2: 3-5 Years Experience
doc2 = Document()
add_heading(doc2, "DAVID SANTOS", 0)
add_paragraph(
    doc2,
    "📍 Austin, TX | ✉️ david.santos@gmail.com | 📞 (512) 555-7890 | 🔗 linkedin.com/in/davidsantosdev | github.com/davidsdev",
)
add_heading(doc2, "SUMMARY", 1)
add_paragraph(
    doc2,
    "Results-oriented Backend Engineer with 4 years of experience designing, developing, and deploying production-grade backend services...",
)
add_heading(doc2, "TECHNICAL SKILLS", 1)
add_paragraph(
    doc2,
    "- Languages: Python, Go, JavaScript (Node.js)\n- Frameworks: Django, FastAPI, Express.js\n- Databases: PostgreSQL, MySQL, Redis\n- Cloud & DevOps: AWS, Docker, Terraform, GitHub Actions\n- Other: Kafka, RabbitMQ, Prometheus, Grafana",
)
add_heading(doc2, "PROFESSIONAL EXPERIENCE", 1)
add_paragraph(
    doc2,
    "Backend Engineer II, Fintract Global, Jul 2021 – Present\n- Designed and implemented event-driven microservices...\n",
)
add_paragraph(
    doc2,
    "Backend Engineer I, HealthStack Technologies, Aug 2019 – Jun 2021\n- Built RESTful APIs in Django...",
)
add_heading(doc2, "PROJECTS", 1)
add_paragraph(doc2, "OpenScribe - An open-source API for speech-to-text processing...")
add_heading(doc2, "EDUCATION & CERTIFICATION", 1)
add_paragraph(
    doc2,
    "B.S. in Software Engineering, University of Texas at Austin – 2019\nAWS Certified Developer – Associate, 2023",
)
resumes.append(("Backend_Engineer_3-5_Years.docx", doc2))

# Resume 3: 5+ Years Experience
doc3 = Document()
add_heading(doc3, "FRANK LIU", 0)
add_paragraph(
    doc3,
    "📍 New York, NY | ✉️ frank.liu@devmail.com | 📞 (917) 321-4567 | 🔗 linkedin.com/in/frankliu | github.com/frankliudev | Portfolio: frankliu.dev",
)
add_heading(doc3, "SUMMARY", 1)
add_paragraph(
    doc3,
    "Senior Backend Engineer with 6+ years of experience architecting and scaling backend systems across fintech and enterprise SaaS domains...",
)
add_heading(doc3, "CORE COMPETENCIES", 1)
add_paragraph(
    doc3,
    "- Languages: Python, Java, Go\n- Architectures: Microservices, Event-driven, Serverless\n- Cloud: AWS, GCP\n- Databases: PostgreSQL, DynamoDB, Cassandra\n- Tools: Kubernetes, Docker, Terraform, Kafka\n- Security: OAuth2, RBAC, API Gateway\n- Leadership: Code reviews, Tech planning",
)
add_heading(doc3, "PROFESSIONAL EXPERIENCE", 1)
add_paragraph(
    doc3,
    "Senior Backend Engineer, Accenture, Feb 2022 – Present\n- Led backend development for financial APIs...\n",
)
add_paragraph(
    doc3,
    "Backend Engineer, Nova Systems, Jul 2018 – Jan 2022\n- Architected data ingestion pipelines...",
)
add_heading(doc3, "PROJECTS & INITIATIVES", 1)
add_paragraph(
    doc3,
    "CostGuard Initiative - Optimized AWS utilization saving ~$100K annually\nDevMentor Circle - Co-founded internal mentorship program",
)
add_heading(doc3, "EDUCATION & CERTIFICATIONS", 1)
add_paragraph(
    doc3,
    "B.S. Computer Science, Northeastern University – 2018\nCertified Kubernetes Application Developer (CKAD)\nAWS Certified Solutions Architect – Associate",
)
resumes.append(("Backend_Engineer_5+_Years.docx", doc3))

# Save files
paths = []
for filename, document in resumes:
    path = f"/mnt/data/{filename}"
    document.save(path)
    paths.append(path)

paths
